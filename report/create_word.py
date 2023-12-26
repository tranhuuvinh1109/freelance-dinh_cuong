import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def add_paragraph_with_style(doc, text, size, is_bold, alignment):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(size)
    font.bold = is_bold
    paragraph.alignment = alignment
    return paragraph
def create_report_docx(filename,save_dir, report):
    file_path = os.path.join(save_dir, filename)
    doc = Document()
    date_object = datetime.strptime(report["end_day"], "%d/%m/%Y")

    add_paragraph_with_style(doc, f'Báo cáo ADG Trạm {report["location"]} ngày {report["start_day"]} đến ngày {report["end_day"]}', 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER)

    # Thêm dòng trắng
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'Kính gửi: Lãnh đạo Đài VT QNN!', 16, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'Trạm {report["location"]} báo cáo tình hình thông tin liên lạc từ 07h00’ ngày {report["start_day"]} đến 07h00’ ngày {report["end_day"]}.', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, 'I. Tình hình thông tin:', 16, True, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'1. Thiết bị viễn thông: {report["device"]}', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'2. Cáp quang: {report["cable"]}', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'3. Nguồn điện, điều hoà: {report["power"]}', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, 'II. Tình hình công việc:', 16, True, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, '1. Thực hiện theo công văn:', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'{report["report"]}', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, '2. Công việc khác:', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'{report["other_job"]}', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, 'III. Tồn tại:', 16, True, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'{report["exist"]}', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, 'ĐỀ XUẤT KIẾN NGHỊ', 16, True, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'{report["propose"]}', 12, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph_with_style(doc, f'Quy Nhơn, Ngày {date_object.day} tháng {date_object.month} năm {date_object.year}', 12, False, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    add_paragraph_with_style(doc, 'Người báo cáo', 12, False, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    add_paragraph_with_style(doc, f'{report["creator"]}', 14, True, WD_PARAGRAPH_ALIGNMENT.RIGHT)

    doc.save(file_path)

def clear_word_document(file_path):
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            paragraph.clear()

        doc.save(file_path)

        return {'message': 'Document cleared successfully'}

    except Exception as e:
        return {'message': f"An error occurred: {str(e)}"}
